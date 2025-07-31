import os
import tempfile
from pathlib import Path
from urllib.parse import urlparse
from typing import Union

from crewai_tools.rag.loaders.base_loader import BaseLoader, LoaderResult


class DOCXLoader(BaseLoader):
    def load(self, source: Union[str, Path], **kwargs) -> LoaderResult:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX loading. Install with: 'pip install python-docx' or pip install crewai-tools[rag]")

        source_str = str(source)

        if self._is_url(source):
            temp_file = self._download_from_url(source_str, kwargs)
            try:
                return self._load_from_file(temp_file, source_str, DocxDocument)
            finally:
                os.unlink(temp_file)
        elif os.path.exists(source_str):
            return self._load_from_file(source_str, source_str, DocxDocument)
        else:
            raise ValueError(f"Source must be a valid file path or URL, got: {source}")

    def _is_url(self, source: Union[str, Path]) -> bool:
        if not isinstance(source, str):
            return False
        try:
            parsed_url = urlparse(source)
            return bool(parsed_url.scheme and parsed_url.netloc)
        except Exception:
            return False

    def _download_from_url(self, url: str, kwargs: dict) -> str:
        import requests

        headers = kwargs.get("headers", {
            "Accept": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "User-Agent": "Mozilla/5.0 (compatible; crewai-tools DOCXLoader)"
        })

        try:
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()

            # Create temporary file to save the DOCX content
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(response.content)
                return temp_file.name
        except Exception as e:
            raise ValueError(f"Error fetching DOCX from URL {url}: {str(e)}")

    def _load_from_file(self, file_path: str, source_str: str, DocxDocument) -> LoaderResult:
        try:
            doc = DocxDocument(file_path)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            content = "\n".join(text_parts)

            metadata = {
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }

            return LoaderResult(content=content, source=source_str, metadata=metadata)

        except Exception as e:
            raise ValueError(f"Error loading DOCX file: {str(e)}")
